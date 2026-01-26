#!/usr/bin/env python3
"""Generate DOCX documentation for GATRA Comparison Framework."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

def set_cell_shading(cell, color):
    """Set cell background color."""
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_with_number(doc, text, level):
    """Add a heading with proper formatting."""
    heading = doc.add_heading(text, level=level)
    return heading

def create_table(doc, headers, rows, col_widths=None):
    """Create a formatted table."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], "D9E2F3")

    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, cell_data in enumerate(row_data):
            row.cells[i].text = str(cell_data)

    return table

def generate_technical_documentation():
    """Generate Technical Documentation for Developers."""
    doc = Document()

    # Title
    title = doc.add_heading('GATRA Comparison Framework', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Technical Documentation for Developers')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f'Version 1.0 | January 2026')
    doc.add_page_break()

    # Table of Contents placeholder
    doc.add_heading('Table of Contents', 1)
    toc_items = [
        '1. Overview',
        '2. Architecture',
        '3. Algorithm Implementation',
        '4. Data Pipeline',
        '5. Configuration',
        '6. API Reference',
        '7. Performance Characteristics',
        '8. Known Issues & Limitations',
    ]
    for item in toc_items:
        doc.add_paragraph(item)
    doc.add_page_break()

    # 1. Overview
    doc.add_heading('1. Overview', 1)
    doc.add_paragraph(
        'The GATRA Comparison Framework provides a rigorous comparison between two anomaly detection '
        'architectures for Security Operations Center (SOC) threat detection:'
    )

    doc.add_paragraph('• Architecture A: PPO (Proximal Policy Optimization) - Reinforcement Learning approach', style='List Bullet')
    doc.add_paragraph('• Architecture B: Isolation Forest - Traditional ML baseline', style='List Bullet')

    doc.add_heading('1.1 Key Features', 2)
    features = [
        'Fair comparison with identical data splits and threshold tuning',
        'Separate latency measurements per model',
        'Fixed calibration for consistent Isolation Forest scoring',
        'Configurable reward shaping for RL policy',
        'Comprehensive evaluation metrics (F1, ROC-AUC, PR-AUC)',
    ]
    for f in features:
        doc.add_paragraph(f'• {f}', style='List Bullet')

    doc.add_heading('1.2 Project Structure', 2)
    structure = '''
gatra-comparison/
├── architecture_a_rl/       # PPO implementation
│   ├── networks.py          # Actor-Critic neural networks
│   ├── ppo.py               # PPO update algorithm
│   ├── env_bandit.py        # Reward function (contextual bandit)
│   └── train.py             # Training loop with validation
├── architecture_b_iforest/  # Isolation Forest implementation
│   ├── model.py             # IForestModel with calibration
│   ├── preprocess.py        # StandardScaler preprocessor
│   └── train.py             # Training with threshold tuning
├── data/                    # Data loading and processing
│   ├── schemas.py           # Pydantic models (RawEvent, Label)
│   ├── loaders.py           # Data loading utilities
│   ├── splits.py            # Time-based train/val/test splits
│   ├── features.py          # Feature extraction (v7, v128)
│   └── toy.py               # Synthetic dataset generator
├── evaluation/              # Evaluation utilities
│   ├── metrics.py           # Classification metrics
│   └── head_to_head.py      # Comparative evaluation
├── configs/                 # YAML configuration files
│   ├── data.yaml            # Dataset configuration
│   ├── ppo.yaml             # PPO hyperparameters
│   ├── iforest.yaml         # Isolation Forest parameters
│   └── eval.yaml            # Evaluation settings
└── tests/                   # Unit tests
'''
    doc.add_paragraph(structure, style='No Spacing')

    # 2. Architecture
    doc.add_heading('2. Architecture', 1)

    doc.add_heading('2.1 PPO (Reinforcement Learning)', 2)
    doc.add_paragraph(
        'The PPO implementation uses a contextual bandit formulation rather than full RL. '
        'This is a deliberate design decision appropriate for SOC alert triage where:'
    )
    doc.add_paragraph('• Each security event is treated independently (no sequential dependencies)', style='List Bullet')
    doc.add_paragraph('• Rewards are immediate based on action and ground truth label', style='List Bullet')
    doc.add_paragraph('• Actions are sampled once before training (offline/batch RL)', style='List Bullet')

    doc.add_heading('Neural Network Architecture', 3)
    doc.add_paragraph('Actor Network: 128 → 256 → 128 → 64 → 4 (softmax)')
    doc.add_paragraph('Critic Network: 128 → 256 → 128 → 64 → 1')
    doc.add_paragraph('Total Parameters: 148,677')

    doc.add_heading('Actions', 3)
    actions_table = [
        ['escalate', 'Escalate to senior analyst', 'High-priority threats'],
        ['contain', 'Initiate containment', 'Active threats requiring response'],
        ['monitor', 'Continue monitoring', 'Suspicious but uncertain'],
        ['dismiss', 'Mark as benign', 'False positives'],
    ]
    create_table(doc, ['Action', 'Description', 'Use Case'], actions_table)

    doc.add_heading('2.2 Isolation Forest', 2)
    doc.add_paragraph(
        'Standard sklearn IsolationForest with fixed calibration from training set statistics. '
        'Key parameters:'
    )
    if_params = [
        ['n_estimators', '200', 'Number of isolation trees'],
        ['max_samples', 'auto', 'Samples per tree'],
        ['contamination', 'auto', 'Expected anomaly ratio'],
        ['max_features', '1.0', 'Features per tree'],
    ]
    create_table(doc, ['Parameter', 'Default', 'Description'], if_params)

    # 3. Algorithm Implementation
    doc.add_heading('3. Algorithm Implementation', 1)

    doc.add_heading('3.1 PPO Update', 2)
    doc.add_paragraph('The PPO update implements the clipped surrogate objective:')
    doc.add_paragraph('L_CLIP = E[min(r_t * A_t, clip(r_t, 1-ε, 1+ε) * A_t)]')
    doc.add_paragraph('Where r_t = π_new(a|s) / π_old(a|s) and ε = 0.2 (clip ratio)')

    doc.add_heading('Key Implementation Details', 3)
    doc.add_paragraph('• Advantage normalization: (A - mean) / (std + 1e-8)', style='List Bullet')
    doc.add_paragraph('• Entropy bonus coefficient: 0.01', style='List Bullet')
    doc.add_paragraph('• Value loss coefficient: 0.5', style='List Bullet')
    doc.add_paragraph('• Gradient clipping: max_norm = 0.5', style='List Bullet')

    doc.add_heading('3.2 Reward Function', 2)
    doc.add_paragraph('Reward shaping for the contextual bandit:')
    reward_table = [
        ['True Positive', 'threat + escalate/contain', 'tp_base × severity', '+10.0 × sev'],
        ['False Negative', 'threat + monitor/dismiss', 'fn_base × severity', '-15.0 × sev'],
        ['True Negative', 'benign + monitor/dismiss', 'efficiency_bonus', '+1.0'],
        ['False Positive', 'benign + escalate/contain', 'fp_base × action_cost', '-3.0 × cost'],
    ]
    create_table(doc, ['Outcome', 'Condition', 'Formula', 'Example'], reward_table)

    doc.add_heading('3.3 Isolation Forest Scoring', 2)
    doc.add_paragraph(
        'The IForestModel uses fixed calibration from training set to ensure consistent scoring:'
    )
    doc.add_paragraph('1. Fit calibration: Store min/max raw scores from training data')
    doc.add_paragraph('2. Score normalization: score = (max - raw) / (max - min)')
    doc.add_paragraph('3. Clip to [0, 1] range')
    doc.add_paragraph('This ensures single-sample scoring equals batch scoring.')

    # 4. Data Pipeline
    doc.add_heading('4. Data Pipeline', 1)

    doc.add_heading('4.1 Data Schemas', 2)
    doc.add_paragraph('RawEvent fields:')
    event_fields = [
        ['event_id', 'str', 'Unique event identifier'],
        ['ts', 'datetime', 'Event timestamp'],
        ['src_ip', 'str | None', 'Source IP address'],
        ['dst_ip', 'str | None', 'Destination IP address'],
        ['port', 'int | None', 'Destination port'],
        ['protocol', 'str | None', 'Protocol (tcp/udp)'],
        ['duration', 'float | None', 'Connection duration'],
        ['bytes_sent', 'float | None', 'Bytes sent'],
        ['bytes_received', 'float | None', 'Bytes received'],
    ]
    create_table(doc, ['Field', 'Type', 'Description'], event_fields)

    doc.add_heading('4.2 Feature Extraction', 2)
    doc.add_paragraph('v7 Features (Isolation Forest):')
    v7_features = [
        ['0', 'duration', 'Connection duration'],
        ['1', 'bytes_sent', 'Bytes sent'],
        ['2', 'bytes_received', 'Bytes received'],
        ['3', 'port', 'Destination port'],
        ['4', 'protocol_encoded', '1.0 if TCP, else 0.0'],
        ['5', 'hour', 'Hour of day (0-23)'],
        ['6', 'dow', 'Day of week (0-6)'],
    ]
    create_table(doc, ['Index', 'Feature', 'Description'], v7_features)

    doc.add_paragraph()
    doc.add_paragraph('v128 Features (RL): Includes v7 features plus cyclic time encodings. Note: Only 11 of 128 dimensions are currently populated.')

    doc.add_heading('4.3 Time-Based Splits', 2)
    splits_table = [
        ['Train', '2025-01-01', '2025-08-31', '~66%'],
        ['Validation', '2025-09-01', '2025-10-31', '~17%'],
        ['Test', '2025-11-01', '2025-12-31', '~17%'],
    ]
    create_table(doc, ['Split', 'Start', 'End', 'Proportion'], splits_table)

    # 5. Configuration
    doc.add_heading('5. Configuration', 1)

    doc.add_heading('5.1 PPO Configuration (configs/ppo.yaml)', 2)
    ppo_config = '''
rl:
  algo: "ppo"
  seed: 42
  state_dim: 128
  action_dim: 4

reward:
  tp_base: 10.0
  fp_base: -3.0
  fn_base: -15.0
  efficiency_bonus: 1.0
  action_cost:
    escalate: 3.0
    contain: 2.0
    monitor: 1.0
    dismiss: 0.5

networks:
  hidden_sizes: [256, 128, 64]

train:
  epochs: 10
  batch_size: 64
  lr: 0.0003
  clip_ratio: 0.2
  entropy_coef: 0.01
  value_coef: 0.5
  max_grad_norm: 0.5

validation:
  eval_every_epoch: 1
  early_stop_patience: 3
'''
    doc.add_paragraph(ppo_config, style='No Spacing')

    doc.add_heading('5.2 Isolation Forest Configuration (configs/iforest.yaml)', 2)
    if_config = '''
model:
  name: "isolation_forest"
  random_state: 42
  n_estimators: 200
  max_samples: "auto"
  contamination: "auto"
  max_features: 1.0
  bootstrap: false

scoring:
  threshold: 0.8
'''
    doc.add_paragraph(if_config, style='No Spacing')

    # 6. API Reference
    doc.add_heading('6. API Reference', 1)

    doc.add_heading('6.1 Training Functions', 2)
    doc.add_paragraph('train_ppo(ppo_cfg_path: str, data_cfg_path: str) -> str')
    doc.add_paragraph('    Trains PPO policy and returns path to saved checkpoint.')
    doc.add_paragraph()
    doc.add_paragraph('train_iforest(iforest_cfg: str, data_cfg: str) -> str')
    doc.add_paragraph('    Trains Isolation Forest and returns path to saved bundle.')

    doc.add_heading('6.2 Evaluation Functions', 2)
    doc.add_paragraph('run_head_to_head(eval_cfg, data_cfg, iforest_cfg, ppo_cfg) -> str')
    doc.add_paragraph('    Runs comparative evaluation and returns path to report JSON.')
    doc.add_paragraph()
    doc.add_paragraph('classification_metrics(y_true, y_score, threshold) -> dict')
    doc.add_paragraph('    Computes precision, recall, F1, ROC-AUC, PR-AUC.')

    # 7. Performance
    doc.add_heading('7. Performance Characteristics', 1)

    doc.add_heading('7.1 Training Performance', 2)
    training_perf = [
        ['Isolation Forest', '~113 ms', '3,319 events'],
        ['PPO (10 epochs)', '~926 ms', '3,319 events'],
    ]
    create_table(doc, ['Model', 'Training Time', 'Dataset Size'], training_perf)

    doc.add_heading('7.2 Inference Latency', 2)
    inference_perf = [
        ['Single-sample', '5.09 ms', '0.03 ms'],
        ['Batch (100)', '0.06 ms/sample', '0.001 ms/sample'],
        ['Batch (823)', '0.013 ms/sample', '0.0004 ms/sample'],
    ]
    create_table(doc, ['Mode', 'Isolation Forest', 'PPO (RL)'], inference_perf)

    doc.add_heading('7.3 Memory Usage', 2)
    memory_table = [
        ['Model parameters', '~2.5 MB', '580 KB'],
        ['Feature storage (train)', '91 KB', '1,660 KB'],
    ]
    create_table(doc, ['Component', 'Isolation Forest', 'PPO (RL)'], memory_table)

    # 8. Known Issues
    doc.add_heading('8. Known Issues & Limitations', 1)

    doc.add_heading('8.1 Algorithm Issues', 2)
    issues = [
        ['Medium', 'FP penalty asymmetry', 'Action cost applied multiplicatively to FP but not TP'],
        ['Medium', 'Sparse v128 features', '117/128 dimensions are zeros'],
        ['Low', 'Action cost ignored for TP', 'Escalating threats has no cost penalty'],
        ['Low', 'Toy data leakage', 'Threats have strong signals (ports, bytes)'],
    ]
    create_table(doc, ['Priority', 'Issue', 'Description'], issues)

    doc.add_heading('8.2 Performance Issues', 2)
    perf_issues = [
        ['High', 'IForest single-sample latency', '5ms/event too slow for real-time; use batching'],
        ['Medium', 'PPO data preparation', '29% of training time; could be vectorized'],
    ]
    create_table(doc, ['Priority', 'Issue', 'Recommendation'], perf_issues)

    # Save
    os.makedirs('docs', exist_ok=True)
    doc.save('docs/01_Technical_Documentation.docx')
    print('Generated: docs/01_Technical_Documentation.docx')


def generate_support_guide():
    """Generate Technical Support Guide."""
    doc = Document()

    # Title
    title = doc.add_heading('GATRA Comparison Framework', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Technical Support Guide')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Version 1.0 | January 2026')
    doc.add_page_break()

    # TOC
    doc.add_heading('Table of Contents', 1)
    toc = [
        '1. Quick Start',
        '2. Installation & Setup',
        '3. Running the System',
        '4. Troubleshooting Guide',
        '5. Common Error Messages',
        '6. FAQ',
        '7. Support Escalation',
    ]
    for item in toc:
        doc.add_paragraph(item)
    doc.add_page_break()

    # 1. Quick Start
    doc.add_heading('1. Quick Start', 1)
    doc.add_paragraph('Minimum steps to run the comparison framework:')

    quick_start = '''
# 1. Create virtual environment
python3 -m venv .venv
source .venv/bin/activate

# 2. Install dependencies
pip install numpy pandas torch scikit-learn pyyaml pydantic joblib

# 3. Run tests to verify installation
python -m pytest tests/ -v

# 4. Train both models
python -m architecture_b_iforest.train
python -m architecture_a_rl.train

# 5. Run evaluation
python -m evaluation.head_to_head
'''
    doc.add_paragraph(quick_start, style='No Spacing')

    # 2. Installation
    doc.add_heading('2. Installation & Setup', 1)

    doc.add_heading('2.1 System Requirements', 2)
    requirements = [
        ['Python', '3.10+', '3.11 recommended'],
        ['RAM', '4 GB minimum', '8 GB recommended'],
        ['Disk', '500 MB', 'For models and data'],
        ['OS', 'Linux/macOS/Windows', 'All supported'],
    ]
    create_table(doc, ['Component', 'Minimum', 'Recommended'], requirements)

    doc.add_heading('2.2 Dependencies', 2)
    deps = [
        ['numpy', '>=1.24.0', 'Numerical operations'],
        ['pandas', '>=2.0.0', 'Data manipulation'],
        ['torch', '>=2.0.0', 'Neural networks (CPU)'],
        ['scikit-learn', '>=1.3.0', 'Isolation Forest, metrics'],
        ['pyyaml', '>=6.0', 'Configuration files'],
        ['pydantic', '>=2.0', 'Data validation'],
        ['joblib', '>=1.3.0', 'Model serialization'],
    ]
    create_table(doc, ['Package', 'Version', 'Purpose'], deps)

    doc.add_heading('2.3 Installation Steps', 2)
    doc.add_paragraph('Option A: Using pip (Recommended)')
    doc.add_paragraph('pip install -r requirements.txt', style='No Spacing')
    doc.add_paragraph()
    doc.add_paragraph('Option B: Manual installation')
    doc.add_paragraph('pip install numpy pandas torch scikit-learn pyyaml pydantic joblib', style='No Spacing')

    # 3. Running
    doc.add_heading('3. Running the System', 1)

    doc.add_heading('3.1 Training Models', 2)
    doc.add_paragraph('Train Isolation Forest:')
    doc.add_paragraph('python -c "from architecture_b_iforest.train import train_iforest; train_iforest(\'configs/iforest.yaml\', \'configs/data.yaml\')"', style='No Spacing')
    doc.add_paragraph()
    doc.add_paragraph('Train PPO:')
    doc.add_paragraph('python -c "from architecture_a_rl.train import train_ppo; train_ppo(\'configs/ppo.yaml\', \'configs/data.yaml\')"', style='No Spacing')

    doc.add_heading('3.2 Running Evaluation', 2)
    doc.add_paragraph('Run head-to-head comparison:')
    doc.add_paragraph('python -c "from evaluation.head_to_head import run_head_to_head; run_head_to_head(\'configs/eval.yaml\', \'configs/data.yaml\', \'configs/iforest.yaml\', \'configs/ppo.yaml\')"', style='No Spacing')

    doc.add_heading('3.3 Output Files', 2)
    outputs = [
        ['artifacts/iforest/iforest_bundle.joblib', 'Trained IF model + preprocessor'],
        ['artifacts/iforest/train_log.json', 'IF training metrics'],
        ['artifacts/ppo/ppo_policy.pt', 'Trained PPO actor/critic'],
        ['artifacts/ppo/train_log.json', 'PPO training metrics'],
        ['reports/head_to_head_report.json', 'Comparison results'],
    ]
    create_table(doc, ['File', 'Description'], outputs)

    # 4. Troubleshooting
    doc.add_heading('4. Troubleshooting Guide', 1)

    doc.add_heading('4.1 Installation Issues', 2)

    doc.add_paragraph('Problem: ModuleNotFoundError: No module named \'torch\'')
    doc.add_paragraph('Solution: Install PyTorch: pip install torch')
    doc.add_paragraph()

    doc.add_paragraph('Problem: ImportError: cannot import name \'RawEvent\'')
    doc.add_paragraph('Solution: Ensure you\'re running from the project root directory')
    doc.add_paragraph()

    doc.add_paragraph('Problem: YAML parsing error')
    doc.add_paragraph('Solution: Check YAML file indentation (use spaces, not tabs)')

    doc.add_heading('4.2 Training Issues', 2)

    doc.add_paragraph('Problem: Training loss is NaN')
    doc.add_paragraph('Possible causes:')
    doc.add_paragraph('• Learning rate too high (try 0.0001)', style='List Bullet')
    doc.add_paragraph('• Data contains NaN values', style='List Bullet')
    doc.add_paragraph('Solution: Check data with: df.isnull().sum()')
    doc.add_paragraph()

    doc.add_paragraph('Problem: Model not improving')
    doc.add_paragraph('Possible causes:')
    doc.add_paragraph('• Learning rate too low', style='List Bullet')
    doc.add_paragraph('• Insufficient epochs', style='List Bullet')
    doc.add_paragraph('• Class imbalance too severe', style='List Bullet')

    doc.add_heading('4.3 Inference Issues', 2)

    doc.add_paragraph('Problem: Isolation Forest inference too slow')
    doc.add_paragraph('Cause: Single-sample inference has ~5ms overhead per call')
    doc.add_paragraph('Solution: Batch events together (100+ samples) for 400x speedup')
    doc.add_paragraph()

    doc.add_paragraph('Problem: Inconsistent IF scores between runs')
    doc.add_paragraph('Cause: Calibration not using fixed training statistics')
    doc.add_paragraph('Solution: Ensure fit_calibration() was called during training')

    # 5. Error Messages
    doc.add_heading('5. Common Error Messages', 1)

    errors = [
        ['FileNotFoundError: configs/data.yaml', 'Run from project root directory'],
        ['KeyError: \'dataset\'', 'Check data.yaml has dataset section'],
        ['RuntimeError: CUDA out of memory', 'Reduce batch_size in ppo.yaml'],
        ['ValueError: Input contains NaN', 'Clean data or check feature extraction'],
        ['AssertionError in ppo_update', 'Check batch tensor shapes'],
    ]
    create_table(doc, ['Error Message', 'Solution'], errors)

    doc.add_heading('5.1 Feature Validation Error Codes', 2)
    feature_errors = [
        ['FEATURE_DIM_MISMATCH', 'Wrong feature length (v7=7, v128=128)', 'Verify feature version and schema alignment'],
        ['FEATURE_NAN_INF', 'NaN or Inf detected in features', 'Inspect upstream feature pipeline for invalid math or missing values'],
        ['FEATURE_DTYPE', 'Non-numeric feature values', 'Ensure features are numeric floats (no strings/objects)'],
    ]
    create_table(doc, ['Error Code', 'Meaning', 'What To Do'], feature_errors)

    # 6. FAQ
    doc.add_heading('6. Frequently Asked Questions', 1)

    faqs = [
        ('Q: Can I use GPU for training?',
         'A: Yes, but the default config uses CPU. For GPU, add device="cuda" to tensor operations.'),
        ('Q: How do I add my own data?',
         'A: Implement a new loader in data/loaders.py following the CSV/Parquet stub patterns.'),
        ('Q: What threshold should I use?',
         'A: Enable tune_threshold in configs to automatically find optimal F1 threshold on validation set.'),
        ('Q: How do I change the number of actions?',
         'A: Modify action_dim in ppo.yaml and update ACTIONS list in env_bandit.py.'),
        ('Q: Why is RL using 128 features but only 11 are populated?',
         'A: This is a placeholder for future feature expansion (sliding windows, entropy, etc.).'),
    ]
    for q, a in faqs:
        doc.add_paragraph(q, style='List Bullet')
        doc.add_paragraph(a)
        doc.add_paragraph()

    # 7. Support Escalation
    doc.add_heading('7. Support Escalation', 1)

    doc.add_heading('7.1 Escalation Levels', 2)
    escalation = [
        ['L1', 'Installation, configuration', 'This guide'],
        ['L2', 'Training failures, data issues', 'Senior engineer'],
        ['L3', 'Algorithm bugs, performance', 'ML team lead'],
    ]
    create_table(doc, ['Level', 'Scope', 'Contact'], escalation)

    doc.add_heading('7.2 Information to Collect', 2)
    doc.add_paragraph('When escalating, provide:')
    doc.add_paragraph('• Python version: python --version', style='List Bullet')
    doc.add_paragraph('• Package versions: pip freeze', style='List Bullet')
    doc.add_paragraph('• Full error traceback', style='List Bullet')
    doc.add_paragraph('• Config files used', style='List Bullet')
    doc.add_paragraph('• Steps to reproduce', style='List Bullet')

    # Save
    doc.save('docs/02_Technical_Support_Guide.docx')
    print('Generated: docs/02_Technical_Support_Guide.docx')


def generate_operations_guide():
    """Generate Operations & Maintenance Guide."""
    doc = Document()

    # Title
    title = doc.add_heading('GATRA Comparison Framework', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Operations & Maintenance Guide')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Version 1.0 | January 2026')
    doc.add_page_break()

    # TOC
    doc.add_heading('Table of Contents', 1)
    toc = [
        '1. Deployment Overview',
        '2. System Requirements',
        '3. Deployment Procedures',
        '4. Monitoring & Health Checks',
        '5. Maintenance Procedures',
        '6. Backup & Recovery',
        '7. Performance Tuning',
        '8. Security Considerations',
        '9. Runbooks',
    ]
    for item in toc:
        doc.add_paragraph(item)
    doc.add_page_break()

    # 1. Deployment Overview
    doc.add_heading('1. Deployment Overview', 1)

    doc.add_heading('1.1 System Components', 2)
    components = [
        ['Training Service', 'Trains models on new data', 'Batch (daily/weekly)'],
        ['Inference Service', 'Real-time threat scoring', 'Always-on'],
        ['Evaluation Service', 'Compares model performance', 'On-demand'],
        ['Data Pipeline', 'Ingests and preprocesses events', 'Streaming/Batch'],
    ]
    create_table(doc, ['Component', 'Purpose', 'Schedule'], components)

    doc.add_heading('1.2 Deployment Architecture', 2)
    doc.add_paragraph('''
┌─────────────────┐     ┌─────────────────┐
│   SIEM/EDR      │────▶│  Data Pipeline  │
└─────────────────┘     └────────┬────────┘
                                 │
                    ┌────────────┴────────────┐
                    ▼                         ▼
           ┌────────────────┐       ┌────────────────┐
           │  IF Inference  │       │  RL Inference  │
           └───────┬────────┘       └───────┬────────┘
                   │                         │
                   └──────────┬──────────────┘
                              ▼
                    ┌─────────────────┐
                    │   Ensemble /    │
                    │   Alert Router  │
                    └─────────────────┘
''', style='No Spacing')

    # 2. System Requirements
    doc.add_heading('2. System Requirements', 1)

    doc.add_heading('2.1 Hardware Requirements', 2)
    hw_req = [
        ['Development', '4 cores', '8 GB', '50 GB SSD'],
        ['Production (low)', '8 cores', '16 GB', '100 GB SSD'],
        ['Production (high)', '16 cores', '32 GB', '500 GB SSD'],
    ]
    create_table(doc, ['Environment', 'CPU', 'RAM', 'Storage'], hw_req)

    doc.add_heading('2.2 Throughput Capacity', 2)
    throughput = [
        ['Single-sample (real-time)', '196 EPS', '32,258 EPS'],
        ['Batched (batch=100)', '16,667 EPS', '1,000,000 EPS'],
        ['Recommended mode', 'Batched', 'Either'],
    ]
    create_table(doc, ['Mode', 'Isolation Forest', 'PPO (RL)'], throughput)

    doc.add_paragraph()
    doc.add_paragraph('EPS = Events Per Second')
    doc.add_paragraph('Note: Isolation Forest requires batching for production throughput.')

    # 3. Deployment Procedures
    doc.add_heading('3. Deployment Procedures', 1)

    doc.add_heading('3.1 Initial Deployment', 2)
    doc.add_paragraph('Step 1: Environment Setup')
    doc.add_paragraph('''
# Create production directory
mkdir -p /opt/gatra
cd /opt/gatra

# Clone or copy codebase
cp -r /path/to/gatra-comparison/* .

# Create virtual environment
python3 -m venv .venv
source .venv/bin/activate

# Install dependencies
pip install -r requirements.txt
''', style='No Spacing')

    doc.add_paragraph()
    doc.add_paragraph('Step 2: Configuration')
    doc.add_paragraph('''
# Copy and edit configurations
cp configs/data.yaml configs/data_prod.yaml
cp configs/ppo.yaml configs/ppo_prod.yaml
cp configs/iforest.yaml configs/iforest_prod.yaml

# Edit paths and parameters for production
# - Set output_dir to /var/lib/gatra/models
# - Adjust batch_size based on available memory
# - Set appropriate thresholds
''', style='No Spacing')

    doc.add_paragraph()
    doc.add_paragraph('Step 3: Initial Training')
    doc.add_paragraph('''
# Train both models
python -c "from architecture_b_iforest.train import train_iforest; \\
    train_iforest('configs/iforest_prod.yaml', 'configs/data_prod.yaml')"

python -c "from architecture_a_rl.train import train_ppo; \\
    train_ppo('configs/ppo_prod.yaml', 'configs/data_prod.yaml')"

# Verify models were created
ls -la /var/lib/gatra/models/
''', style='No Spacing')

    doc.add_heading('3.2 Model Update Deployment', 2)
    doc.add_paragraph('Blue-Green deployment for zero-downtime updates:')
    doc.add_paragraph('''
# 1. Train new models to staging directory
export STAGING_DIR=/var/lib/gatra/models_staging

# 2. Run validation
python -m pytest tests/ -v

# 3. Run evaluation on holdout set
python -c "from evaluation.head_to_head import run_head_to_head; ..."

# 4. Compare metrics with production
# If new model is better:

# 5. Atomic swap
mv /var/lib/gatra/models /var/lib/gatra/models_old
mv /var/lib/gatra/models_staging /var/lib/gatra/models

# 6. Restart inference service
systemctl restart gatra-inference

# 7. Verify health
curl http://localhost:8080/health
''', style='No Spacing')

    # 4. Monitoring
    doc.add_heading('4. Monitoring & Health Checks', 1)

    doc.add_heading('4.1 Key Metrics to Monitor', 2)
    metrics = [
        ['Inference Latency', 'P95 < 10ms (batched)', 'Prometheus/Grafana'],
        ['Throughput', '> 1000 EPS', 'Application logs'],
        ['Error Rate', '< 0.1%', 'Application logs'],
        ['Invalid Feature Rate', '< 0.5%', 'Application logs/metrics'],
        ['Model Drift', 'F1 > 0.7', 'Weekly evaluation'],
        ['Memory Usage', '< 80% of limit', 'System metrics'],
        ['CPU Usage', '< 70% average', 'System metrics'],
    ]
    create_table(doc, ['Metric', 'Target', 'Source'], metrics)

    doc.add_heading('4.2 Health Check Endpoints', 2)
    doc.add_paragraph('Recommended health check implementation:')
    health_checks = [
        ['/health', 'Basic liveness', '200 OK'],
        ['/health/ready', 'Model loaded', '200 OK + model version'],
        ['/health/live', 'Can process events', '200 OK + latency'],
    ]
    create_table(doc, ['Endpoint', 'Purpose', 'Expected Response'], health_checks)

    doc.add_heading('4.3 Alerting Rules', 2)
    alerts = [
        ['Critical', 'Inference error rate > 5%', 'Page on-call'],
        ['Critical', 'No events processed for 5min', 'Page on-call'],
        ['Warning', 'Latency P95 > 50ms', 'Slack notification'],
        ['Warning', 'invalid_feature_total rate spike', 'Slack notification + investigate'],
        ['Warning', 'Memory > 80%', 'Slack notification'],
        ['Info', 'Model retrained', 'Email notification'],
    ]
    create_table(doc, ['Severity', 'Condition', 'Action'], alerts)

    # 5. Maintenance
    doc.add_heading('5. Maintenance Procedures', 1)

    doc.add_heading('5.1 Scheduled Maintenance', 2)
    maintenance = [
        ['Daily', 'Log rotation', 'Automated (logrotate)'],
        ['Weekly', 'Model evaluation', 'Run head_to_head on new data'],
        ['Monthly', 'Model retraining', 'Train on last 3 months data'],
        ['Quarterly', 'Dependency updates', 'pip install --upgrade'],
        ['Annually', 'Full system review', 'Architecture review'],
    ]
    create_table(doc, ['Frequency', 'Task', 'Notes'], maintenance)

    doc.add_heading('5.2 Log Management', 2)
    doc.add_paragraph('Log locations:')
    logs = [
        ['/var/log/gatra/inference.log', 'Inference service logs'],
        ['/var/log/gatra/training.log', 'Training job logs'],
        ['/var/log/gatra/evaluation.log', 'Evaluation results'],
    ]
    create_table(doc, ['Path', 'Contents'], logs)

    doc.add_paragraph()
    doc.add_paragraph('Log retention policy:')
    doc.add_paragraph('• Inference logs: 7 days', style='List Bullet')
    doc.add_paragraph('• Training logs: 90 days', style='List Bullet')
    doc.add_paragraph('• Evaluation reports: 1 year', style='List Bullet')

    doc.add_heading('5.3 Model Lifecycle', 2)
    doc.add_paragraph('''
Model Version Naming: gatra-{model}-{date}-{hash}
Example: gatra-ppo-20260115-a1b2c3d

Retention Policy:
- Keep last 3 production models
- Keep all models from last 30 days
- Archive older models to cold storage
''', style='No Spacing')

    # 6. Backup & Recovery
    doc.add_heading('6. Backup & Recovery', 1)

    doc.add_heading('6.1 What to Backup', 2)
    backups = [
        ['Model files', '/var/lib/gatra/models/', 'After each training'],
        ['Configuration', '/opt/gatra/configs/', 'After changes'],
        ['Training data', '/var/lib/gatra/data/', 'Daily'],
        ['Evaluation reports', '/var/lib/gatra/reports/', 'Weekly'],
    ]
    create_table(doc, ['Component', 'Location', 'Frequency'], backups)

    doc.add_heading('6.2 Backup Procedure', 2)
    doc.add_paragraph('''
#!/bin/bash
# backup_gatra.sh

BACKUP_DIR=/backup/gatra/$(date +%Y%m%d)
mkdir -p $BACKUP_DIR

# Backup models
cp -r /var/lib/gatra/models $BACKUP_DIR/

# Backup configs
cp -r /opt/gatra/configs $BACKUP_DIR/

# Compress
tar -czf $BACKUP_DIR.tar.gz $BACKUP_DIR

# Upload to S3 (optional)
aws s3 cp $BACKUP_DIR.tar.gz s3://company-backups/gatra/
''', style='No Spacing')

    doc.add_heading('6.3 Recovery Procedure', 2)
    doc.add_paragraph('''
#!/bin/bash
# restore_gatra.sh BACKUP_DATE

BACKUP_DATE=$1
BACKUP_FILE=/backup/gatra/${BACKUP_DATE}.tar.gz

# Stop services
systemctl stop gatra-inference

# Extract backup
tar -xzf $BACKUP_FILE -C /tmp/

# Restore models
cp -r /tmp/${BACKUP_DATE}/models/* /var/lib/gatra/models/

# Restore configs
cp -r /tmp/${BACKUP_DATE}/configs/* /opt/gatra/configs/

# Restart services
systemctl start gatra-inference

# Verify health
sleep 5
curl http://localhost:8080/health
''', style='No Spacing')

    # 7. Performance Tuning
    doc.add_heading('7. Performance Tuning', 1)

    doc.add_heading('7.1 Isolation Forest Tuning', 2)
    if_tuning = [
        ['n_estimators', '200', '50-100', 'Reduces latency 2-4x'],
        ['Batch size', '1', '100+', 'Critical: 400x speedup'],
        ['n_jobs', '-1', '-1', 'Use all cores'],
    ]
    create_table(doc, ['Parameter', 'Default', 'Tuned', 'Impact'], if_tuning)

    doc.add_heading('7.2 PPO Tuning', 2)
    ppo_tuning = [
        ['batch_size', '64', '128-256', 'Better GPU utilization'],
        ['hidden_sizes', '[256,128,64]', '[128,64]', 'Faster inference'],
        ['state_dim', '128', '11', 'Match actual features'],
    ]
    create_table(doc, ['Parameter', 'Default', 'Tuned', 'Impact'], ppo_tuning)

    doc.add_heading('7.3 System Tuning', 2)
    doc.add_paragraph('Linux kernel parameters for high throughput:')
    doc.add_paragraph('''
# /etc/sysctl.conf additions
net.core.somaxconn = 65535
net.ipv4.tcp_max_syn_backlog = 65535
vm.swappiness = 10
''', style='No Spacing')

    # 8. Security
    doc.add_heading('8. Security Considerations', 1)

    doc.add_heading('8.1 Access Control', 2)
    doc.add_paragraph('• Model files: Read-only for inference service', style='List Bullet')
    doc.add_paragraph('• Config files: Read-only, no secrets in configs', style='List Bullet')
    doc.add_paragraph('• Training jobs: Run as dedicated service account', style='List Bullet')
    doc.add_paragraph('• API endpoints: Require authentication', style='List Bullet')

    doc.add_heading('8.2 Data Protection', 2)
    doc.add_paragraph('• Event data may contain sensitive information', style='List Bullet')
    doc.add_paragraph('• Encrypt data at rest and in transit', style='List Bullet')
    doc.add_paragraph('• Implement data retention policies', style='List Bullet')
    doc.add_paragraph('• Log access to model predictions', style='List Bullet')

    doc.add_heading('8.3 Model Security', 2)
    doc.add_paragraph('• Sign model files to prevent tampering', style='List Bullet')
    doc.add_paragraph('• Validate model checksums before loading', style='List Bullet')
    doc.add_paragraph('• Monitor for adversarial inputs', style='List Bullet')

    # 9. Runbooks
    doc.add_heading('9. Runbooks', 1)

    doc.add_heading('9.1 Runbook: High Latency Alert', 2)
    doc.add_paragraph('''
Trigger: Inference latency P95 > 50ms

Steps:
1. Check current load
   - Look at events/second in metrics
   - Compare to baseline

2. Check batch size
   - If processing single events, enable batching
   - Target batch size: 100+

3. Check resource usage
   - CPU: Should be < 80%
   - Memory: Should be < 80%

4. If IF model is slow:
   - Consider reducing n_estimators (50-100)
   - Ensure batching is enabled

5. If RL model is slow:
   - Check for memory pressure
   - Consider reducing hidden layer sizes

6. Escalate if not resolved in 15 minutes
''', style='No Spacing')

    doc.add_heading('9.2 Runbook: Model Drift Detected', 2)
    doc.add_paragraph('''
Trigger: Weekly evaluation shows F1 < 0.7

Steps:
1. Verify evaluation data quality
   - Check for label errors
   - Verify data pipeline integrity

2. Analyze drift
   - Compare feature distributions
   - Check for new attack patterns

3. If data quality OK:
   - Schedule emergency retraining
   - Use last 30 days of labeled data

4. After retraining:
   - Run full evaluation suite
   - Compare with production model
   - Deploy if metrics improve

5. Document incident
   - Root cause
   - New patterns detected
   - Model version deployed
''', style='No Spacing')

    doc.add_heading('9.3 Runbook: Service Restart', 2)
    doc.add_paragraph('''
When: After deployment, config change, or recovery

Steps:
1. Notify stakeholders
   - Post in #soc-alerts channel
   - Expected downtime: < 1 minute

2. Stop service gracefully
   systemctl stop gatra-inference

3. Verify stopped
   systemctl status gatra-inference

4. Clear any stale state
   rm -f /tmp/gatra-*.lock

5. Start service
   systemctl start gatra-inference

6. Verify health
   curl http://localhost:8080/health
   curl http://localhost:8080/health/ready

7. Monitor for 5 minutes
   - Check error rates
   - Check latency
   - Check throughput

8. Notify completion
   - Post in #soc-alerts channel
''', style='No Spacing')

    doc.add_heading('9.4 Runbook: Invalid Feature Rate Spike', 2)
    doc.add_paragraph('''
Trigger: invalid_feature_total rate increases above baseline

Steps:
1. Inspect recent requests for bad payloads
   - Sample error responses and request IDs
   - Identify feature_version (v7/v128)

2. Verify upstream feature generator
   - Check for schema mismatch (dims)
   - Check for NaN/Inf sources (division by zero, missing data)
   - Check for non-numeric fields in ingestion

3. If caused by client regression
   - Roll back client or feature pipeline version
   - Add validation at producer side

4. If unknown
   - Increase logging around request parsing
   - Escalate to data pipeline owner
''', style='No Spacing')

    # Save
    doc.save('docs/03_Operations_Maintenance_Guide.docx')
    print('Generated: docs/03_Operations_Maintenance_Guide.docx')


if __name__ == '__main__':
    generate_technical_documentation()
    generate_support_guide()
    generate_operations_guide()
    print('\nAll documents generated successfully!')
